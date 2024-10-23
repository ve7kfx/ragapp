import sys
import os
import shutil
import argparse
import numpy as np
import uuid
import datetime
import json
import warnings
import faiss  # Faiss library for similarity search
from langchain_chroma import Chroma
from langchain_community.embeddings import GPT4AllEmbeddings, OllamaEmbeddings
from langchain.prompts import PromptTemplate, ChatPromptTemplate
from langchain_ollama import OllamaLLM  # Updated import
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain.chains import RetrievalQA
from langchain.schema import Document
from tkinter import Tk, filedialog
from pdfminer.high_level import extract_text
from docx import Document as DocxDocument  # Added for .docx support
from bs4 import BeautifulSoup  # Added for HTML parsing
import requests  # Added for URL handling
from rich.console import Console
from rich.prompt import Prompt, Confirm
from rich.table import Table
from rich.progress import Progress, SpinnerColumn, BarColumn, TextColumn
import questionary
import zipfile  # Added for zip functionality
import tempfile  # Added for temporary directory management
import newspaper  # Added for advanced web scraping

# Initialize Rich Console
console = Console()

# Suppress specific Pydantic warnings (temporary fix)
warnings.filterwarnings(
    "ignore",
    category=UserWarning,
    message="Field \"model_name\" in GPT4AllEmbeddings has conflict with protected namespace \"model_\"."
)

# Paths for Chroma and data storage
CHROMA_PATH = "chroma"
DATA_PATH = "data"
FAISS_INDEX_PATH = "faiss_index.bin"  # Changed from hnsw_index.bin
ID_MAPPING_PATH = "id_mapping.json"
CONFIG_PATH = "config.json"
MEMORY_LOG_PATH = "memory_log.json"  # Path for memory logging

# Define the default system prompt
DEFAULT_SYSTEM_PROMPT = (
    "You are an intelligent assistant designed to answer questions based on the provided context. "
    "Use your reasoning abilities to generate comprehensive and accurate responses."
)

# Set the correct embedding dimension based on the model
EMBEDDING_DIM = 768  # Update this based on your embedding model's output

# Initialize Faiss index (CPU only)
def initialize_faiss_index(embedding_dim):
    try:
        index_flat = faiss.IndexFlatIP(embedding_dim)  # Inner Product (for cosine similarity)
        faiss_index = faiss.IndexIDMap(index_flat)  # Wrap in IndexIDMap to add custom IDs
        console.print("[green]Faiss CPU index initialized.[/green]")
        return faiss_index
    except Exception as e:
        console.print(f"[red]Error initializing Faiss index: {e}[/red]")
        sys.exit(1)

# Placeholder for Faiss index (will be initialized later)
faiss_index = None

# Custom Embeddings Classes to Resolve Pydantic Warning
class SafeGPT4AllEmbeddings(GPT4AllEmbeddings):
    class Config:
        protected_namespaces = ()

class SafeOllamaEmbeddings(OllamaEmbeddings):
    class Config:
        protected_namespaces = ()

# Config management functions
def load_config():
    """Load the configuration from config.json. If it doesn't exist, create it with default values."""
    if os.path.exists(CONFIG_PATH):
        try:
            with open(CONFIG_PATH, "r") as f:
                config = json.load(f)
                system_prompt = config.get("system_prompt", DEFAULT_SYSTEM_PROMPT)
                console.print(f"Loaded system prompt from {CONFIG_PATH}.")
                return system_prompt
        except Exception as e:
            console.print(f"[red]Error loading config: {e}[/red]")
            return DEFAULT_SYSTEM_PROMPT
    else:
        # If config doesn't exist, create it with the default system prompt
        try:
            with open(CONFIG_PATH, "w") as f:
                json.dump({"system_prompt": DEFAULT_SYSTEM_PROMPT}, f, indent=4)
            console.print(f"[yellow]No config found. Created {CONFIG_PATH} with default system prompt.[/yellow]")
        except Exception as e:
            console.print(f"[red]Error creating config file: {e}[/red]")
        return DEFAULT_SYSTEM_PROMPT

def save_config(system_prompt):
    """Save the system prompt to config.json."""
    try:
        with open(CONFIG_PATH, "w") as f:
            json.dump({"system_prompt": system_prompt}, f, indent=4)
        console.print(f"[green]System prompt saved to {CONFIG_PATH}.[/green]")
    except Exception as e:
        console.print(f"[red]Error saving config: {e}[/red]")

# Function to clear the memory log
def clear_memory_log():
    if os.path.exists(MEMORY_LOG_PATH):
        try:
            os.remove(MEMORY_LOG_PATH)
            console.print(f"[green]Memory log '{MEMORY_LOG_PATH}' cleared.[/green]")
        except Exception as e:
            console.print(f"[red]Error clearing memory log: {e}[/red]")
    else:
        console.print(f"[yellow]Memory log '{MEMORY_LOG_PATH}' does not exist.[/yellow]")

# Embedding function using OllamaLLM
def get_embedding_function(embedding_model="nomic-embed-text"):
    console.print(f"Initializing embedding function with '{embedding_model}'...")
    try:
        if embedding_model == "nomic-embed-text":
            embeddings = SafeOllamaEmbeddings(model="nomic-embed-text:latest")
        elif embedding_model == "nomic-embed-code":
            embeddings = SafeOllamaEmbeddings(model="nomic-embed-code:latest")
        elif embedding_model == "another-embed-model":
            embeddings = SafeOllamaEmbeddings(model="another-embed-model:latest")
        else:
            console.print(f"[yellow]Embedding model '{embedding_model}' is not recognized. Using default 'nomic-embed-text'.[/yellow]")
            embeddings = SafeOllamaEmbeddings(model="nomic-embed-text:latest")
        return embeddings
    except Exception as e:
        console.print(f"[red]Error initializing OllamaEmbeddings: {e}[/red]")
        sys.exit(1)

# Clear the entire database including Chroma, Faiss index, and mappings
def clear_database():
    paths_to_clear = [CHROMA_PATH, FAISS_INDEX_PATH, ID_MAPPING_PATH, CONFIG_PATH, MEMORY_LOG_PATH]
    for path in paths_to_clear:
        if os.path.isdir(path):
            shutil.rmtree(path)
            console.print(f"[green]Directory {path} cleared.[/green]")
        elif os.path.isfile(path):
            os.remove(path)
            console.print(f"[green]File {path} deleted.[/green]")
    # Reinitialize config with default prompt
    load_config()

# Load documents supporting multiple formats and URLs
def load_documents(input_dir=None, urls=None, temp_dir=None):
    documents = []
    if input_dir:
        if not os.path.exists(input_dir):
            console.print(f"[red]Input directory {input_dir} does not exist.[/red]")
        else:
            with Progress(SpinnerColumn(), TextColumn("[progress.description]{task.description}")) as progress:
                task = progress.add_task("Loading documents from directory...", total=None)
                for root, _, files in os.walk(input_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        try:
                            if file.lower().endswith(".pdf"):
                                text = extract_text(file_path)
                            elif file.lower().endswith(".txt"):
                                with open(file_path, "r", encoding="utf-8") as f:
                                    text = f.read()
                            elif file.lower().endswith(".docx"):
                                doc = DocxDocument(file_path)
                                text = "\n".join([para.text for para in doc.paragraphs])
                            elif file.lower().endswith(".html") or file.lower().endswith(".htm"):
                                with open(file_path, "r", encoding="utf-8") as f:
                                    soup = BeautifulSoup(f, "html.parser")
                                    text = soup.get_text()
                            else:
                                console.print(f"[yellow]Unsupported file type: {file_path}. Skipping.[/yellow]")
                                continue

                            if text.strip():
                                doc = Document(page_content=text, metadata={"source": file})
                                documents.append(doc)
                                console.print(f"[green]Loaded {file_path} successfully.[/green]")
                            else:
                                console.print(f"[yellow]No text found in {file_path}. Skipping.[/yellow]")
                        except Exception as e:
                            console.print(f"[red]Error loading {file_path}: {e}[/red]")
    if urls:
        with Progress(SpinnerColumn(), TextColumn("[progress.description]{task.description}")) as progress:
            task = progress.add_task("Loading documents from URLs...", total=len(urls))
            for url in urls:
                try:
                    paper = newspaper.Article(url)
                    paper.download()
                    paper.parse()
                    text = paper.text
                    if text.strip():
                        if temp_dir:
                            # Generate a unique filename
                            filename = uuid.uuid4().hex + ".txt"
                            temp_file_path = os.path.join(temp_dir, filename)
                            with open(temp_file_path, "w", encoding="utf-8") as f:
                                f.write(text)
                            # Read the text back from the temp file
                            with open(temp_file_path, "r", encoding="utf-8") as f:
                                temp_text = f.read()
                            doc = Document(page_content=temp_text, metadata={"source": url})
                            documents.append(doc)
                            console.print(f"[green]Loaded content from {url} and saved to temporary file.[/green]")
                        else:
                            doc = Document(page_content=text, metadata={"source": url})
                            documents.append(doc)
                            console.print(f"[green]Loaded content from {url} successfully.[/green]")
                    else:
                        console.print(f"[yellow]No text found at {url}. Skipping.[/yellow]")
                except Exception as e:
                    console.print(f"[red]Error scraping {url}: {e}[/red]")
                progress.advance(task)
    if not documents:
        console.print("[yellow]No valid documents found.[/yellow]")
    return documents

# Split documents into chunks
def split_documents(documents):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=80,
        length_function=len
    )
    chunks = text_splitter.split_documents(documents)
    console.print(f"[green]Split into {len(chunks)} chunks.[/green]")
    return chunks

# Calculate chunk IDs for each chunk
def calculate_chunk_ids(chunks):
    """Calculate chunk IDs based on the document's source and page."""
    last_page_id = None
    current_chunk_index = 0

    for chunk in chunks:
        source = chunk.metadata.get("source", "unknown")
        # Assuming each Document corresponds to one page; adjust if multiple pages per Document
        page = chunk.metadata.get("page", 0)
        current_page_id = f"{source}:{page}"

        # If the page ID is the same as the last one, increment the index.
        if current_page_id == last_page_id:
            current_chunk_index += 1
        else:
            current_chunk_index = 0

        # Calculate the chunk ID.
        chunk_id = f"{current_page_id}:{current_chunk_index}"
        last_page_id = current_page_id

        # Add it to the chunk's metadata.
        chunk.metadata["id"] = chunk_id

    return chunks

# Load ID mapping
def load_id_mapping():
    if os.path.exists(ID_MAPPING_PATH):
        try:
            with open(ID_MAPPING_PATH, "r") as f:
                mapping = json.load(f)
                console.print(f"[green]Loaded ID mapping with {len(mapping)} entries.[/green]")
                return mapping
        except Exception as e:
            console.print(f"[red]Error loading ID mapping: {e}[/red]")
            return {}
    console.print("[yellow]No existing ID mapping found. Starting fresh.[/yellow]")
    return {}

# Save ID mapping
def save_id_mapping(mapping):
    try:
        with open(ID_MAPPING_PATH, "w") as f:
            json.dump(mapping, f)
        console.print("[green]ID mapping saved successfully.[/green]")
    except Exception as e:
        console.print(f"[red]Error saving ID mapping: {e}[/red]")

# Function to save Faiss index
def save_faiss_index(faiss_index, path):
    try:
        # Save the index directly as we are using CPU only
        console.print("[yellow]Saving index directly...[/yellow]")
        faiss.write_index(faiss_index, path)
        console.print(f"[green]Faiss index saved to {path}.[/green]")
    except Exception as e:
        console.print(f"[red]Error saving Faiss index: {e}[/red]")

# Load Faiss index from the file system
def load_faiss_index():
    global faiss_index
    if os.path.exists(FAISS_INDEX_PATH):
        try:
            # Load the underlying index
            faiss_index = faiss.read_index(FAISS_INDEX_PATH)
            console.print("[green]Faiss index loaded successfully.[/green]")
        except Exception as e:
            console.print(f"[red]Error loading Faiss index: {e}[/red]")
            faiss_index = initialize_faiss_index(EMBEDDING_DIM)
    else:
        faiss_index = initialize_faiss_index(EMBEDDING_DIM)
        console.print("[yellow]No existing Faiss index found. A new index has been initialized.[/yellow]")

# Add document embeddings to Chroma and Faiss
def add_to_faiss_and_chroma(db, chunks, id_mapping, embedding_function):
    global faiss_index
    # Calculate unique chunk IDs
    chunks_with_ids = calculate_chunk_ids(chunks)

    # Fetch existing document IDs from Chroma
    try:
        existing_items = db.get(include=[])  # IDs are always included by default
        existing_ids = set(existing_items["ids"])
        console.print(f"[green]Number of existing documents in Chroma DB: {len(existing_ids)}.[/green]")
    except Exception as e:
        console.print(f"[red]Error fetching existing document IDs from Chroma: {e}[/red]")
        existing_ids = set()

    # Filter new chunks
    new_chunks = [chunk for chunk in chunks_with_ids if chunk.metadata["id"] not in existing_ids]

    if new_chunks:
        console.print(f"[green]Adding {len(new_chunks)} new documents.[/green]")
        new_chunk_ids = [chunk.metadata["id"] for chunk in new_chunks]

        # Add to Chroma
        try:
            db.add_documents(new_chunks, ids=new_chunk_ids)
            console.print("[green]Documents added to Chroma.[/green]")
        except Exception as e:
            console.print(f"[red]Error adding documents to Chroma: {e}[/red]")
            return

        # Embed documents
        try:
            with Progress(SpinnerColumn(), TextColumn("[progress.description]{task.description}")) as progress:
                task = progress.add_task("Embedding documents...", total=len(new_chunks))
                embeddings = []
                for chunk in new_chunks:
                    emb = embedding_function.embed_documents([chunk.page_content])[0]
                    embeddings.append(emb)
                    progress.advance(task)
            console.print("[green]Documents embedded.[/green]")
        except Exception as e:
            console.print(f"[red]Error embedding documents: {e}[/red]")
            return

        # Debug: Check embeddings
        for idx, emb in enumerate(embeddings[:3]):  # Inspect first 3 embeddings
            sample = emb[:5] if isinstance(emb, (list, tuple, np.ndarray)) else "N/A"
            console.print(f"Embedding {idx}: Type: {type(emb)}, Length: {len(emb)}, Sample: {sample}")

        # Ensure embeddings are lists of floats
        processed_embeddings = []
        for emb in embeddings:
            if isinstance(emb, (list, tuple, np.ndarray)):
                if all(isinstance(x, float) for x in emb):
                    processed_embeddings.append(emb)
                else:
                    try:
                        processed_emb = [float(x) for x in emb]
                        processed_embeddings.append(processed_emb)
                    except Exception as e:
                        console.print(f"[red]Error converting embedding to floats: {e}[/red]")
                        processed_embeddings.append([0.0] * EMBEDDING_DIM)  # Fallback to zeros
            else:
                console.print(f"[yellow]Unexpected embedding type: {type(emb)}. Using zeros.[/yellow]")
                processed_embeddings.append([0.0] * EMBEDDING_DIM)  # Fallback to zeros

        # Update the ID mapping
        for chunk_id in new_chunk_ids:
            if chunk_id not in id_mapping:
                id_mapping[chunk_id] = len(id_mapping) + 1  # Assign a new integer ID

        # Save the updated mapping
        save_id_mapping(id_mapping)

        # Prepare embeddings for Faiss
        try:
            embeddings_array = np.array(processed_embeddings).astype('float32')
            faiss.normalize_L2(embeddings_array)  # Normalize for cosine similarity
        except Exception as e:
            console.print(f"[red]Error processing embeddings: {e}[/red]")
            embeddings_array = np.zeros((len(processed_embeddings), EMBEDDING_DIM), dtype='float32')

        # Assign unique integer IDs to embeddings
        integer_ids = [id_mapping[chunk.metadata["id"]] for chunk in new_chunks]
        integer_ids_array = np.array(integer_ids, dtype=np.int64)

        # Add embeddings to Faiss index
        try:
            faiss_index.add_with_ids(embeddings_array, integer_ids_array)
            console.print("[green]New documents added to Faiss index.[/green]")
        except Exception as e:
            console.print(f"[red]Error adding embeddings to Faiss: {e}[/red]")
            return
    else:
        console.print("[yellow]No new documents to add.[/yellow]")

    # Save the Faiss index
    try:
        save_faiss_index(faiss_index, FAISS_INDEX_PATH)
    except Exception as e:
        console.print(f"[red]Error saving Faiss index: {e}[/red]")

# Query function for RAG system with LLM Agent Prompt
def query_rag(query_text, model_name, db, embedding_function, system_prompt, id_mapping):
    # Embed the query
    try:
        query_embedding = embedding_function.embed_documents([query_text])[0]
        console.print("[green]Query embedded.[/green]")
    except Exception as e:
        console.print(f"[red]Error embedding the query: {e}[/red]")
        query_embedding = [0.0] * EMBEDDING_DIM  # Fallback

    # Ensure query_embedding is a numpy array of floats and normalize
    try:
        query_embedding = np.array(query_embedding, dtype=np.float32).reshape(1, -1)
        faiss.normalize_L2(query_embedding)
    except Exception as e:
        console.print(f"[red]Error processing query embedding: {e}[/red]")
        query_embedding = np.array([[0.0] * EMBEDDING_DIM], dtype=np.float32)  # Fallback

    # Perform similarity search using Faiss
    try:
        with Progress(SpinnerColumn(), TextColumn("[progress.description]{task.description}")) as progress:
            task = progress.add_task("Performing similarity search...", total=None)
            distances, indices = faiss_index.search(query_embedding, k=5)
            progress.advance(task)
        console.print("[green]Similarity search completed.[/green]")
    except Exception as e:
        console.print(f"[red]Error performing similarity search: {e}[/red]")
        distances, indices = np.array([]), np.array([])

    # Retrieve documents based on indices
    results = []
    if indices.size > 0 and len(id_mapping) > 0:
        for idx in indices[0]:
            # Reverse lookup in ID mapping
            chunk_id = None
            for key, value in id_mapping.items():
                if value == idx:
                    chunk_id = key
                    break
            if chunk_id:
                # Fetch document from Chroma
                try:
                    result = db.get(ids=[chunk_id])
                    if result['documents']:
                        doc_text = result['documents'][0]
                        results.append(Document(page_content=doc_text, metadata={"id": chunk_id}))
                    else:
                        console.print(f"[yellow]No document found with ID {chunk_id} in Chroma.[/yellow]")
                except Exception as e:
                    console.print(f"[red]Error retrieving document with ID {chunk_id} from Chroma: {e}[/red]")

    if results:
        context_text = "\n\n---\n\n".join([f"Document {doc.metadata['id']}: {doc.page_content}" for doc in results])
    else:
        context_text = "No relevant documents found."

    # **Enhanced Prompt Template for LLM Agent with System Prompt**
    prompt_template = ChatPromptTemplate.from_template(f"""
    {system_prompt}

    **Context:**
    {{context}}

    **Question:**
    {{question}}

    **Answer:**
    """)
    prompt = prompt_template.format(context=context_text, question=query_text)

    # Invoke the LLM
    try:
        model = OllamaLLM(model=model_name)  # Updated class
        response_text = model.invoke(prompt)
        console.print("[green]LLM invoked successfully.[/green]")
    except Exception as e:
        console.print(f"[red]Error invoking the LLM: {e}[/red]")
        response_text = "I'm sorry, I couldn't process your request."

    sources = [doc.metadata['id'] for doc in results]
    formatted_response = f"Response: {response_text}\nSources: {sources}"

    # Log the user query and response to memory log
    log_memory(query_text, response_text)

    console.print(formatted_response)
    return formatted_response

# Log chat interactions
def log_chat(user_input, response):
    timestamp = datetime.datetime.now().isoformat()
    log_entry = f"{timestamp}: User Input: {user_input}, Response: {response}"
    try:
        with open("chat_log.txt", "a") as log_file:
            log_file.write(log_entry + "\n")
        console.print(f"[green]Chat logged at {timestamp}.[/green]")
    except Exception as e:
        console.print(f"[red]Error logging chat: {e}[/red]")

# Log memory interactions to memory_log.json
def log_memory(prompt, response):
    timestamp = datetime.datetime.now().isoformat()
    entry = {
        "timestamp": timestamp,
        "prompt": prompt,
        "response": response
    }
    try:
        if os.path.exists(MEMORY_LOG_PATH):
            with open(MEMORY_LOG_PATH, "a") as mem_file:
                mem_file.write(json.dumps(entry) + "\n")
        else:
            with open(MEMORY_LOG_PATH, "w") as mem_file:
                mem_file.write(json.dumps(entry) + "\n")
        console.print(f"[green]Memory logged at {timestamp}.[/green]")
    except Exception as e:
        console.print(f"[red]Error logging memory: {e}[/red]")

# GUI to select the local PDF file
def select_pdf():
    root = Tk()
    root.withdraw()  # Hide the root window
    file_path = filedialog.askopenfilename(
        filetypes=[
            ("PDF files", "*.pdf"),
            ("Text files", "*.txt"),
            ("Word documents", "*.docx"),
            ("HTML files", "*.html;*.htm")
        ]
    )
    root.destroy()
    return file_path

# GUI to input URLs
def input_urls():
    urls = []
    while True:
        url = questionary.text("Enter a URL (or leave blank to finish):").ask()
        if url:
            urls.append(url)
        else:
            break
    return urls

# Embed documents function
def embed_documents(db, id_mapping, embedding_function):
    console.print("[cyan]Embed Command Initiated.[/cyan]")
    choice = questionary.select(
        "Would you like to embed:",
        choices=[
            "Single Document",
            "Directory of Documents",
            "URLs",
            "Cancel"
        ]
    ).ask()

    if choice == "Single Document":
        file_path = select_pdf()
        if not file_path:
            console.print("[yellow]No file selected. Aborting embed operation.[/yellow]")
            return
        if file_path.lower().endswith((".pdf", ".txt", ".docx", ".html", ".htm")):
            try:
                if file_path.lower().endswith(".pdf"):
                    text = extract_text(file_path)
                elif file_path.lower().endswith(".txt"):
                    with open(file_path, "r", encoding="utf-8") as f:
                        text = f.read()
                elif file_path.lower().endswith(".docx"):
                    doc = DocxDocument(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                elif file_path.lower().endswith(".html") or file_path.lower().endswith(".htm"):
                    with open(file_path, "r", encoding="utf-8") as f:
                        soup = BeautifulSoup(f, "html.parser")
                        text = soup.get_text()
                else:
                    console.print(f"[yellow]Unsupported file type: {file_path}. Skipping.[/yellow]")
                    return
                if text.strip():
                    doc = Document(page_content=text, metadata={"source": os.path.basename(file_path)})
                    documents = [doc]
                    console.print(f"[green]Loaded {file_path} successfully.[/green]")
                else:
                    console.print(f"[yellow]No text found in {file_path}. Skipping.[/yellow]")
                    return
            except Exception as e:
                console.print(f"[red]Error loading {file_path}: {e}[/red]")
                return
        else:
            console.print(f"[yellow]Unsupported file type: {file_path}. Skipping.[/yellow]")
            return
    elif choice == "Directory of Documents":
        root = Tk()
        root.withdraw()  # Hide the root window
        directory = filedialog.askdirectory(title="Select Documents Directory")
        root.destroy()
        if not directory:
            console.print("[yellow]No directory selected. Aborting embed operation.[/yellow]")
            return
        documents = load_documents(input_dir=directory, urls=None)
    elif choice == "URLs":
        urls = input_urls()
        if not urls:
            console.print("[yellow]No URLs entered. Aborting embed operation.[/yellow]")
            return
        # Create a temporary directory for saving scraped URL content
        with tempfile.TemporaryDirectory() as temp_dir:
            documents = load_documents(input_dir=None, urls=urls, temp_dir=temp_dir)
            if not documents:
                console.print("[yellow]No documents to embed.[/yellow]")
                return
            console.print(f"[green]Loaded {len(documents)} documents from URLs.[/green]")
            chunks = split_documents(documents)
            add_to_faiss_and_chroma(db, chunks, id_mapping, embedding_function)
            console.print("✅ [green]New documents embedded successfully![/green]")
            # Temporary directory and files are automatically cleaned up here
        return  # Exit the function after handling URLs
    else:
        console.print("[yellow]Embed operation cancelled.[/yellow]")
        return

    if not documents:
        console.print("[yellow]No documents to embed.[/yellow]")
        return

    console.print(f"[green]Loaded {len(documents)} documents.[/green]")
    chunks = split_documents(documents)
    add_to_faiss_and_chroma(db, chunks, id_mapping, embedding_function)
    console.print("✅ [green]New documents embedded successfully![/green]")

# Set System Prompt Function
def set_system_prompt(current_prompt):
    """Allow the user to set a custom system prompt."""
    console.print("[cyan]Set System Prompt[/cyan]")
    new_prompt = questionary.text("Enter your custom system prompt:", default=current_prompt).ask()
    if new_prompt:
        save_config(new_prompt)
        console.print("[green]✅ System prompt updated successfully![/green]")
        return new_prompt
    else:
        console.print("[yellow]No input received. System prompt remains unchanged.[/yellow]")
        return current_prompt

# Save the entire state to a zip file
def save_state(zip_path):
    try:
        # Save Faiss index before zipping
        if hasattr(faiss_index, 'is_trained') and faiss_index.is_trained:
            save_faiss_index(faiss_index, FAISS_INDEX_PATH)

        # Chroma typically handles its own persistence. Ensure the directory is included.
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for folder in [CHROMA_PATH, DATA_PATH]:
                if os.path.isdir(folder):
                    for root_dir, dirs, files in os.walk(folder):
                        for file in files:
                            file_path = os.path.join(root_dir, file)
                            zipf.write(file_path, os.path.relpath(file_path, start=folder))
            for file in [FAISS_INDEX_PATH, ID_MAPPING_PATH, CONFIG_PATH, MEMORY_LOG_PATH]:
                if os.path.isfile(file):
                    zipf.write(file, os.path.basename(file))
        console.print(f"[green]State saved successfully to {zip_path}.[/green]")
    except Exception as e:
        console.print(f"[red]Error saving state: {e}[/red]")

# Load the entire state from a zip file
def load_state(zip_path):
    try:
        with zipfile.ZipFile(zip_path, 'r') as zipf:
            zipf.extractall()
        console.print(f"[green]State loaded successfully from {zip_path}.[/green]")
        # After extracting, load the Faiss index
        load_faiss_index()
    except Exception as e:
        console.print(f"[red]Error loading state: {e}[/red]")

# Display main menu
def display_menu():
    table = Table(title="Main Menu", show_header=True, header_style="bold magenta")
    table.add_column("Option", style="dim", width=20)
    table.add_column("Description")
    table.add_row("1. Embed Documents", "Add new documents (PDF, TXT, DOCX, HTML) or URLs to the database.")
    table.add_row("2. Query", "Ask a question based on the embedded documents.")
    table.add_row("3. Set System Prompt", "Define a custom system prompt for the assistant.")
    table.add_row("4. Save State", "Save the current state to a .zip file.")
    table.add_row("5. Clear Memory Log", "Clear the memory log.")
    table.add_row("6. Help", "Show available commands.")
    table.add_row("7. Exit", "Quit the application.")
    console.print(table)

# Main function with Faiss integration
def main():
    parser = argparse.ArgumentParser(
        description="""\
        This program processes documents (PDF, TXT, DOCX, HTML) and URLs to extract text, converts the text to embeddings using the Ollama model,
        and stores these embeddings in a Chroma vector store. Additionally, it utilizes Faiss for efficient similarity search.
        You can reset the database, specify which Ollama model to use
        for querying and evaluation, choose the input directory containing the documents,
        and select whether to use GPU acceleration for Faiss.
        """
    )

    parser.add_argument(
        "--clear_db",
        action="store_true",
        help="Reset the embeddings database by clearing all stored embeddings."
    )

    parser.add_argument(
        "--clear_memory",
        action="store_true",
        help="Clear the memory log."
    )

    parser.add_argument(
        "--embedding_model",
        type=str,
        choices=["nomic-embed-text", "nomic-embed-code", "another-embed-model"],  # Add more models as needed
        default="nomic-embed-text",
        help="Specify the embedding model to use (default: 'nomic-embed-text')."
    )

    parser.add_argument(
        "--model",
        type=str,
        default="mistral",
        help="Specify the Ollama model to use for querying and evaluation (default: 'mistral')."
    )

    parser.add_argument(
        "--input_dir",
        type=str,
        default="data",
        help="Directory where the input documents are located (default: './data')."
    )

    parser.add_argument(
        "--save_state",
        type=str,
        help="Path to save the current state as a .zip file."
    )

    parser.add_argument(
        "--load_zip",  # Changed from --load_state to --load_zip
        type=str,
        help="Path to load the state from a .zip file."
    )

    args = parser.parse_args()

    # Handle clear_memory if provided
    if args.clear_memory:
        clear_memory_log()

    # Handle load_zip if provided
    if args.load_zip:
        load_state(args.load_zip)

    # Load system_prompt after loading state
    system_prompt = load_config()

    # Get the script's directory (location of the script itself)
    script_dir = os.path.dirname(os.path.realpath(__file__))
    console.print(f"Script directory: {script_dir}")

    # Data directory path relative to the script's directory
    input_dir = os.path.join(script_dir, args.input_dir)

    # Reset the Chroma and Faiss database if --clear_db is provided
    if args.clear_db:
        console.print("[yellow]✨ Clearing Database[/yellow]")
        clear_database()

    # Set the user-specified query/evaluation model
    model_name = args.model

    # Determine the embedding model based on the --embedding_model flag
    embedding_model = args.embedding_model

    # Initialize embedding function and Chroma database BEFORE adding documents
    embedding_function = get_embedding_function(embedding_model=embedding_model)
    try:
        db = Chroma(persist_directory=CHROMA_PATH, embedding_function=embedding_function)
        console.print("[green]Chroma initialized successfully.[/green]")
    except Exception as e:
        console.print(f"[red]Error initializing Chroma: {e}[/red]")
        sys.exit(1)

    # Load ID mapping
    id_mapping = load_id_mapping()

    # Load Faiss index
    load_faiss_index()

    # Initial embedding of documents if input_dir is provided and not clearing the database and not loading from zip
    if not args.clear_db and not args.load_zip:
        documents = load_documents(input_dir=input_dir, urls=None)
        if documents:
            console.print(f"[green]Loaded {len(documents)} documents.[/green]")
            chunks = split_documents(documents)
            add_to_faiss_and_chroma(db, chunks, id_mapping, embedding_function)
            console.print("✅ [green]Documents embedded successfully![/green]")
        else:
            console.print("[yellow]No documents found in the directory![/yellow]")

    # Interactive command loop
    try:
        while True:
            display_menu()
            choice = questionary.select(
                "Select an option:",
                choices=[
                    "Embed Documents",
                    "Query",
                    "Set System Prompt",
                    "Save State",
                    "Clear Memory Log",
                    "Help",
                    "Exit"
                ]
            ).ask()

            if choice == "Exit":
                console.print("[cyan]Goodbye![/cyan]")
                break
            elif choice == "Embed Documents":
                embed_documents(db, id_mapping, embedding_function)
            elif choice == "Query":
                query = questionary.text("Enter your question:").ask()
                if query:
                    response = query_rag(query, model_name, db, embedding_function, system_prompt, id_mapping)
                    log_chat(query, response)
                else:
                    console.print("[yellow]No question entered. Returning to menu.[/yellow]")
            elif choice == "Set System Prompt":
                system_prompt = set_system_prompt(current_prompt=system_prompt)  # Pass current_prompt
            elif choice == "Save State":
                # Ensure Faiss index is saved before zipping
                if hasattr(faiss_index, 'is_trained') and faiss_index.is_trained:
                    save_faiss_index(faiss_index, FAISS_INDEX_PATH)
                zip_path = questionary.text("Enter the path to save the state (e.g., state.zip):").ask()
                if zip_path:
                    save_state(zip_path)
                else:
                    console.print("[yellow]No path entered. Save operation cancelled.[/yellow]")
            elif choice == "Clear Memory Log":
                clear_memory_log()
            elif choice == "Help":
                help_table = Table(title="Help", show_header=False)
                help_table.add_row("Embed Documents", "Add new documents (PDF, TXT, DOCX, HTML) or URLs to the database.")
                help_table.add_row("Query", "Ask a question based on the embedded documents.")
                help_table.add_row("Set System Prompt", "Define a custom system prompt for the assistant.")
                help_table.add_row("Save State", "Save the current state to a .zip file.")
                help_table.add_row("Clear Memory Log", "Clear the memory log.")
                help_table.add_row("Help", "Show available commands and their descriptions.")
                help_table.add_row("Exit", "Quit the application.")
                console.print(help_table)
            else:
                console.print("[red]Invalid choice. Please select a valid option.[/red]")
    finally:
        # Save Faiss index and ID mapping on exit
        if hasattr(faiss_index, 'is_trained') and faiss_index.is_trained:
            save_faiss_index(faiss_index, FAISS_INDEX_PATH)
            console.print("[green]Faiss index saved.[/green]")
        save_id_mapping(id_mapping)
        # Optionally save state on exit
        if args.save_state:
            save_state(args.save_state)

if __name__ == "__main__":
    main()
